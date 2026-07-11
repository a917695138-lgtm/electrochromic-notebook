from pathlib import Path
from zipfile import ZipFile

from docx import Document
from lxml import etree


DOCX = Path(r"D:\Users\ao\Documents\电致变色\实验记录本\outputs\DA凝胶器件综合性能优化实验方案.docx")


def main():
    doc = Document(DOCX)
    print(f"exists={DOCX.exists()} size={DOCX.stat().st_size}")
    print(f"paragraphs={len(doc.paragraphs)} tables={len(doc.tables)}")
    for para in doc.paragraphs:
        if para.text.strip():
            print("first_nonempty=" + para.text.strip())
            break
    section = doc.sections[0]
    print(
        "section_inches="
        f"{section.page_width.inches:.2f},"
        f"{section.page_height.inches:.2f},"
        f"{section.left_margin.inches:.2f},"
        f"{section.right_margin.inches:.2f}"
    )
    with ZipFile(DOCX) as zf:
        xml = zf.read("word/document.xml")
    root = etree.fromstring(xml)
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    texts = "".join(root.xpath(".//w:t/text()", namespaces=ns))
    for key in [
        "DA凝胶稳定性前提下器件综合性能优化方案",
        "D-100",
        "D-5x-rinse",
        "Q_ATO/Q_PBI",
        "Kim",
        "Alesanco",
        "数据记录模板",
    ]:
        print(f"contains:{key}={key in texts}")
    for idx, tbl in enumerate(root.xpath(".//w:tbl", namespaces=ns), 1):
        grid = tbl.xpath("./w:tblGrid/w:gridCol/@w:w", namespaces=ns)
        ind = tbl.xpath("./w:tblPr/w:tblInd/@w:w", namespaces=ns)
        if idx <= 5 or idx > 18:
            total = sum(map(int, grid or [0]))
            print(f"table:{idx}:cols={len(grid)}:sum={total}:ind={(ind or [''])[0]}")


if __name__ == "__main__":
    main()
