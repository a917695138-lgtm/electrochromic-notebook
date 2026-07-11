from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"D:\Users\ao\Documents\电致变色\实验记录本")
OUT = ROOT / "outputs" / "DA凝胶器件综合性能优化实验方案.docx"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(11, 37, 69)
MUTED = RGBColor(89, 89, 89)
HEADER_FILL = "E8EEF5"
LIGHT_FILL = "F4F6F9"
WARN_FILL = "FFF8E8"
GOOD_FILL = "EAF4EE"
BORDER = "C9D3DF"


def set_run_font(run, size=11, bold=None, italic=None, color=None, east_asia="Microsoft YaHei"):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_para_format(paragraph, before=0, after=6, line=1.25, align=None):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    if align is not None:
        paragraph.alignment = align


def add_paragraph(doc, text="", size=11, bold=False, italic=False, color=None, before=0, after=6, line=1.25, align=None):
    p = doc.add_paragraph()
    set_para_format(p, before=before, after=after, line=line, align=align)
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, bold=bold, italic=italic, color=color)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    if level == 1:
        set_para_format(p, before=18, after=10, line=1.25)
        size, color = 16, BLUE
    elif level == 2:
        set_para_format(p, before=14, after=7, line=1.25)
        size, color = 13, BLUE
    else:
        set_para_format(p, before=10, after=5, line=1.25)
        size, color = 12, DARK_BLUE
    r = p.add_run(text)
    set_run_font(r, size=size, bold=True, color=color)
    p.style = f"Heading {min(level, 3)}"
    return p


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text, size=9.2, bold=False, color=None, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    set_para_format(p, before=0, after=0, line=1.15, align=align)
    r = p.add_run(str(text))
    set_run_font(r, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell)


def set_table_borders(table, color=BORDER):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_width(table, widths):
    total = sum(widths)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for w in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Pt(widths[i] / 20)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[i]))
            tc_w.set(qn("w:type"), "dxa")


def add_table(doc, headers, rows, widths, font_size=9.0, header_fill=HEADER_FILL, center_cols=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    set_table_borders(table)
    center_cols = set(center_cols or [])
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, h in enumerate(headers):
        set_cell_shading(hdr.cells[i], header_fill)
        set_cell_text(hdr.cells[i], h, size=9.0, bold=True, color=INK, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.CENTER if i in center_cols else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[i], value, size=font_size, align=align)
    add_paragraph(doc, "", after=3)
    return table


def add_callout(doc, title, body, fill=LIGHT_FILL):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_width(table, [9360])
    set_table_borders(table, color="D5DCE6")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=130, bottom=130, start=180, end=180)
    p = cell.paragraphs[0]
    set_para_format(p, before=0, after=3, line=1.18)
    r = p.add_run(title)
    set_run_font(r, size=10.3, bold=True, color=INK)
    p2 = cell.add_paragraph()
    set_para_format(p2, before=0, after=0, line=1.18)
    r2 = p2.add_run(body)
    set_run_font(r2, size=9.8, color=RGBColor(45, 45, 45))
    add_paragraph(doc, "", after=3)
    return table


def add_meta_line(doc, label, value):
    p = doc.add_paragraph()
    set_para_format(p, before=0, after=2, line=1.15)
    r1 = p.add_run(label + " ")
    set_run_font(r1, size=10.5, bold=True, color=INK)
    r2 = p.add_run(value)
    set_run_font(r2, size=10.5, color=RGBColor(45, 45, 45))


def add_bullet(doc, text):
    p = doc.add_paragraph(style=None)
    set_para_format(p, before=0, after=4, line=1.25)
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    r = p.add_run("• ")
    set_run_font(r, size=10.8, color=INK)
    r2 = p.add_run(text)
    set_run_font(r2, size=10.8)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)

    for style_name, size, color in [
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 12, DARK_BLUE),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
    header_p = section.header.paragraphs[0]
    set_para_format(header_p, after=0, line=1)
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = header_p.add_run("DA凝胶器件优化方案 | 2026-06-29")
    set_run_font(hr, size=8.5, color=MUTED)
    footer_p = section.footer.paragraphs[0]
    set_para_format(footer_p, after=0, line=1)
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer_p.add_run("实验方案表格版")
    set_run_font(fr, size=8.5, color=MUTED)


def build_doc():
    doc = Document()
    configure_document(doc)

    add_paragraph(doc, "实验方案", size=10, bold=True, color=BLUE, before=0, after=2)
    title = add_paragraph(
        doc,
        "DA凝胶稳定性前提下器件综合性能优化方案",
        size=22,
        bold=True,
        color=INK,
        before=0,
        after=4,
        line=1.1,
    )
    subtitle = add_paragraph(
        doc,
        "TiO2 / PBI-Cl / ATO / 器件结构优化的表格化执行版",
        size=12.5,
        color=RGBColor(70, 70, 70),
        after=12,
    )
    add_meta_line(doc, "日期：", "2026-06-29")
    add_meta_line(doc, "适用体系：", "DA（LiTFSI）+ 50% EMIMTFSI 凝胶电解质平台")
    add_meta_line(doc, "核心目的：", "在凝胶稳定性已较好的前提下，优化非凝胶器件结构，实现高对比度、快速响应和长循环稳定性的平衡。")

    add_callout(
        doc,
        "总判断",
        "当前主要矛盾不再是 DA 凝胶是否稳定，而是 PBI-Cl 有效负载、PBI/ATO 电荷匹配、褪色价态完全性与封装/UV 暴露控制。下一步应先做小样本快速诊断，再进入 TiO2 厚度、ATO 退火和面积匹配的精筛。",
        GOOD_FILL,
    )

    add_heading(doc, "1. 一页总览", 1)
    add_table(
        doc,
        ["指标", "优先目标", "最低目标", "备注/分析"],
        [
            ["光学对比度 ΔT", "≥60%", "≥55%", "以 λmax 处 ΔT 为准，同时记录 300-800 nm 最大 ΔT。"],
            ["着色时间 tc", "≤1.2 s", "≤1.5 s", "按达到 90% ΔT 定义，避免采样间隔造成时间失真。"],
            ["褪色时间 tb", "≤0.8 s", "≤1.0 s", "重点检查 -1/-2 价态残留。"],
            ["8000 cycles 保持率", "≥90%", "≥85%", "需记录 1、50、100、300、500 cycles 早期衰减。"],
            ["13000 cycles 保持率", "≥85%", "继续验证", "作为最优条件长循环验证。"],
            ["可重复性", "平行趋势一致", "异常可追溯", "记录染浴体积、面积、膜厚、封装状态和光路位置。"],
        ],
        [1750, 1500, 1500, 4610],
        center_cols=[1, 2],
    )

    add_table(
        doc,
        ["当前瓶颈", "已有证据", "优先动作", "判定结果如何使用"],
        [
            ["PBI-Cl 负载不足或不可控", "100 μL 滴液 ΔT 约 45-50%；过量浸泡 ΔT 约 63-64%。", "D-100 / D-5x / D-10x / D-5x-rinse。", "确认是否进入 TiO2 厚度优化，或先控制吸附/冲洗。"],
            ["高负载与响应时间权衡", "高负载后 tc 从 0.48 s 增至约 1.04 s。", "比较 ΔT 增益和 tc/tb 增幅。", "若 D-10x 增益小而变慢，停止继续提高负载。"],
            ["氧化不完全与残余峰", "PMMA/LiTFSI 需 +0.33 至 +0.43 V 才能氧化回去；循环后 -1 价峰残留。", "测 Q_ATO/Q_PBI；比较 +0.3 V、+0.5 V 褪色。", "决定 ATO 层数、面积比和褪色正压。"],
            ["封装/UV 干扰", "UV 胶和紫外照射会影响变色层或电解质。", "统一遮挡、边缘点胶、封装照片记录。", "若全部条件早期衰减，优先查封装而非继续加 PBI。"],
        ],
        [1900, 2800, 2300, 2360],
    )

    add_heading(doc, "2. 已有实验事实转化", 1)
    add_table(
        doc,
        ["记录日期/体系", "关键结果", "结论与分析", "下一步影响"],
        [
            ["2026-06-07 DA50%E 低负载", "初始 ΔT 约 47.89%；500 cycles 降至 29.68%；5500 cycles 约 30.96；响应 0.48 s。", "低负载响应快但对比度低，衰减主要集中在早期 500 cycles。", "保留为快速响应参照，不作为性能上限路线。"],
            ["2026-06-10 DA50%E 浸泡染浴", "ΔT 约 64%；tc 1.04 s；tb 0.56 s；13000 cycles 保持率约 89.3%。", "过量浸泡显著提升对比度和长循环稳定性，但着色变慢。", "以此作为高对比度基准，继续寻找速度/稳定平衡。"],
            ["2026-06-16 染浴方式对比", "4-5 倍过量浸泡比 100 μL 滴液颜色更深，ΔT 约 63%。", "PBI-Cl 实际吸附量是早期对比度瓶颈。", "优先做染浴体积、覆盖充分性和 DCM 冲洗诊断。"],
            ["PMMA/LiTFSI 对照", "-0.09 V 开始还原，-1.2 V 基本还原；回扫至 +0.43 V -2 价才开始氧化；循环后 -1 价峰残留。", "氧化不完全可能由 ATO 容量不足、面积匹配、电压窗口或 PBI 聚集共同导致。", "不能只调电压，必须做 Q_ATO/Q_PBI 和 ATO 层数诊断。"],
            ["液态器件封装记录", "改良封装后 8000 cycles 保持率约 95.84%；仍观察到电解液下降。", "封装是独立变量，可能掩盖材料本征差异。", "所有优化样品统一 UV 遮挡和边缘密封策略。"],
        ],
        [1650, 2500, 3000, 2210],
        font_size=8.7,
    )

    add_heading(doc, "3. 文献依据到设计变量", 1)
    add_table(
        doc,
        ["文献依据", "可转化结论", "本方案变量", "注意事项"],
        [
            ["Kim et al., 2009, Solar Energy Materials and Solar Cells, DOI: 10.1016/j.solmat.2009.05.007", "TiO2 厚度增加可提高染料负载和 ΔT，但会增加传输路径并拖慢响应。", "Ti-1 / Ti-2 / Ti-3 / Ti-4。", "不追求越厚越好，以 ΔT 增益和 tc 增幅共同判定。"],
            ["Nunes et al., 2020, Applied Sciences, DOI: 10.3390/app10041200", "多孔 TiO2 提供界面面积，但过厚或孔道不连通会限制离子进入。", "记录膜裂纹、均匀性、层数和响应。", "若厚膜 ΔT 增益小、tc 明显变慢，应停止加厚。"],
            ["Lundy et al., 2018, New Journal of Chemistry, DOI: 10.1039/C8NJ04214D", "PBI/TiO2 杂化受表面覆盖度和分子堆积影响。", "D-5x-rinse、冲洗液 UV-vis、循环后残余峰。", "区分有效锚定与物理吸附/聚集。"],
            ["Gratzel, 2003, J. Photochem. Photobiol. C, DOI: 10.1016/S1389-5567(03)00026-1", "染料吸附存在饱和；过长时间或过高浓度可能聚集。", "P-0.5 / P-1.0 / P-1.0-12h / P-0.5x2。", "4-5 倍浸泡已有效，后续要找平台而非盲目加量。"],
            ["Sauvage et al., 2011, JACS, DOI: 10.1021/ja110541t", "吸附温度会影响覆盖度和聚集/复合。", "室温 24 h 与 50 °C 6/12 h 比较。", "升温可缩短时间，但需检查是否引入聚集。"],
            ["Choi et al., 2017, Solar Energy Materials and Solar Cells, DOI: 10.1016/j.solmat.2017.10.001", "离子存储层容量需要与变色层匹配，过低导致反应不完全，过高可能增内阻。", "ATO-1 / ATO-2 / ATO-3，计算 Q_ATO/Q_PBI。", "推荐 Q_ATO/Q_PBI 暂定 1.0-1.5。"],
            ["Lee et al., 2008, Small", "ATO/SnO2 退火影响晶化、导电性和孔结构。", "AT-400 / AT-450 / AT-500。", "550 °C 以上暂不优先，避免影响 FTO 或膜结构。"],
            ["Alesanco et al., 2018, Materials, DOI: 10.3390/ma11030414", "凝胶器件受电极接触、凝胶厚度、封装和面积匹配影响。", "AR-1.0 / AR-1.2 / AR-1.5；统一热压和封装。", "用有效重叠面积而非实际涂布面积来定义面积比。"],
            ["Chen et al., 2020, Scientific Reports；Li et al., 2021, Nature Communications", "响应时间常用达到 90% 光学调制定义。", "光谱采集间隔 0.1-0.2 s。", "统一算法，避免不同实验间响应时间不可比。"],
        ],
        [2300, 2750, 2200, 2110],
        font_size=8.2,
    )

    add_heading(doc, "4. 机制假设与判别实验", 1)
    add_table(
        doc,
        ["假设", "证据", "预测", "判别实验", "若成立则采取"],
        [
            ["H1 PBI-Cl 有效负载决定 ΔT 上限", "100 μL 与 4-5 倍浸泡 ΔT 差异约 15 个百分点。", "负载增加时 PBI 吸收增强、ΔT 增大；平台后 tc 继续变慢。", "D-100 / D-5x / D-10x。", "选平台附近条件进入 TiO2 厚度优化。"],
            ["H2 过度负载/聚集导致响应变慢和残余峰", "高负载 tc 约 1.04 s，循环后残余峰问题存在。", "强 DCM 冲洗或低浓度二次浸泡可降低残余峰。", "D-5x-rinse、PBI 浓度/时间矩阵。", "采用过量浸泡 + 标准化冲洗；必要时降浓度。"],
            ["H3 ATO 容量/面积不足导致氧化不完全", "+0.33 至 +0.43 V 才能氧化回去。", "Q_ATO/Q_PBI <1 时残余峰明显；提高 ATO 后 tb 改善。", "ATO-1/2/3，AR-1.0/1.2/1.5。", "优化 ATO 层数、面积比和褪色电压。"],
            ["H4 封装与 UV 是独立失效源", "UV 照射和 UV 胶会影响变色层/电解质。", "遮挡变色层、限制 UV 胶接触电解质可提升重复性。", "统一封装照片、封装前后颜色和循环记录。", "若所有组早衰，先修封装流程。"],
        ],
        [1500, 2100, 2200, 1800, 1760],
        font_size=8.5,
    )

    add_heading(doc, "5. Phase A 快速诊断矩阵", 1)
    add_callout(
        doc,
        "Phase A 目标",
        "用约 10-14 个器件在最小样品数内判断主控因素。优先级为 PBI 负载量、ATO 容量、电压窗口，封装作为全过程控制变量。",
    )
    add_heading(doc, "A1. PBI-Cl 染浴负载量诊断", 2)
    add_table(
        doc,
        ["组别", "染浴方式", "固定条件", "目的", "器件数", "判据/备注"],
        [
            ["D-100", "100 μL 滴液覆盖", "TiO2 基准；ATO-2；DA50%E；-2.0 V/0 V 与 -2.0 V/+0.3 V。", "低负载对照。", "2", "若 ΔT 明显低但响应快，说明负载不足。"],
            ["D-5x", "4-5 倍体积过量浸泡", "同上。", "已验证高对比度条件。", "2", "作为高对比度基准。"],
            ["D-10x", "完全过量浸泡", "同上。", "判断负载是否还有增益。", "2", "若相对 D-5x ΔT 增益 <3% 且 tc 增加 >30%，停止加量。"],
            ["D-5x-rinse", "4-5 倍浸泡 + 强 DCM 冲洗", "同上。", "区分有效锚定与物理吸附。", "2", "若 ΔT 略低但残余峰和衰减改善，采用标准偏强冲洗。"],
        ],
        [1000, 1850, 2600, 1600, 650, 1660],
        font_size=8.3,
        center_cols=[0, 4],
    )
    add_paragraph(doc, "测试要点：记录染浴体积、浸泡时间、液面是否覆盖 TiO2；拍照；条件允许时测浸泡前后溶液与 DCM 冲洗液 UV-vis；测 CV、CA+光谱、λmax、ΔT、tc/tb；500 cycles 在 1、50、100、300、500 cycles 记录 ΔT 和残余峰。", size=9.6, color=RGBColor(45, 45, 45), after=8)

    add_heading(doc, "A2. ATO 容量诊断", 2)
    add_table(
        doc,
        ["组别", "ATO 层数", "目的", "器件数", "核心测试", "判据"],
        [
            ["ATO-1", "1 层", "低容量/低阻抗对照。", "2", "Q_ATO、Q_PBI、tb、残余峰、500 cycles。", "若 Q_ATO/Q_PBI <1 且残余峰严重，容量不足。"],
            ["ATO-2", "2 层", "当前基准。", "2", "同上。", "作为比较中心。"],
            ["ATO-3", "3 层", "高容量候选。", "2", "同上。", "若 tb 和保持率改善，进入 ATO 优化；若响应变慢明显，避免过量。"],
        ],
        [950, 1100, 1900, 650, 2500, 2260],
        font_size=8.6,
        center_cols=[0, 1, 3],
    )
    add_table(
        doc,
        ["Q_ATO/Q_PBI", "判断", "实验动作"],
        [
            ["<1.0", "对电极容量不足。", "优先增加 ATO 层数或有效面积。"],
            ["1.0-1.5", "推荐区间。", "优先保留。"],
            ["1.5-2.5", "可接受但需观察响应。", "若 tb 变慢或背景电流升高，回调面积/层数。"],
            [">2.5", "容量可能过剩。", "避免无效电流和内阻增加。"],
        ],
        [1900, 3000, 4460],
        font_size=8.8,
        center_cols=[0],
    )

    add_heading(doc, "A3. 褪色电压窗口诊断", 2)
    add_table(
        doc,
        ["条件", "着色电压", "褪色电压", "单步时间", "目的", "判据"],
        [
            ["V0", "-2.0 V", "0 V", "3 s", "温和基准。", "若残余峰明显，说明 0 V 褪色不足。"],
            ["V+0.3", "-2.0 V", "+0.3 V", "3 s", "验证正压褪色。", "若残余峰降低且 500 cycles 不下降，优先采用。"],
            ["V+0.5", "-2.0 V", "+0.5 V", "3 s", "褪色上限。", "若改善有限或衰减加快，避免使用。"],
            ["Vdeep", "-2.5 V", "+0.3 V", "3 s", "更深还原。", "若 ΔT 升高但 100-500 cycles 快速衰减，不用于长循环。"],
        ],
        [900, 1250, 1250, 1000, 1900, 3060],
        font_size=8.8,
        center_cols=[0, 1, 2, 3],
    )

    add_heading(doc, "6. Phase B 针对性优化矩阵", 1)
    add_table(
        doc,
        ["模块", "启动条件", "实验矩阵", "入选/淘汰规则", "备注/分析"],
        [
            ["B1 TiO2 层数", "A1 确认 PBI 负载为主要变量。", "Ti-1、Ti-2、Ti-3、Ti-4，各 2 个器件。", "入选：ΔT ≥60%、tc ≤1.2 s、tb ≤0.8 s、500 cycles ≥85%。Ti-4 若增益 <3% 且 tc 增加 >30%，淘汰。", "建立层数/负载/响应的关系，不把厚膜作为默认答案。"],
            ["B2 ATO 退火", "A2 显示 ATO 对褪色或残余峰影响明显。", "AT-400、AT-450、AT-500，各 2 个器件。", "AT-450 若响应更快且容量不降，优先；AT-500 若透过率下降、膜裂或容量下降，淘汰。", "目标是导电性与孔结构平衡。"],
            ["B3 PBI 浓度/时间", "D-5x 接近平台但残余峰仍明显。", "P-0.5、P-1.0、P-1.0-12h、P-0.5x2。", "比较初始吸收、DCM 冲洗液吸收、ΔT、tc/tb 和循环后 -1 价残余。", "原记录中“1M”按称量与体积应统一为约 1 mM 量级。"],
            ["B4 面积匹配", "固定最优 TiO2/PBI、ATO 和电压窗口后。", "AR-1.0、AR-1.2、AR-1.5。", "若 AR-1.2 优于 AR-1.0，说明略大对电极有利；若 AR-1.5 响应变慢或背景电流增大，淘汰。", "用掩膜或绝缘胶带控制有效重叠面积，减少制备误差。"],
        ],
        [1250, 1900, 1900, 2800, 1510],
        font_size=8.0,
    )

    add_heading(doc, "7. Phase C 复现与长循环", 1)
    add_table(
        doc,
        ["组合", "目的", "平行器件", "必须达到", "备注"],
        [
            ["高对比度优先组合", "验证 ΔT 上限和长循环可保持性。", "3 个", "ΔT ≥60%；8000 cycles ≥90%；13000 cycles ≥85%。", "若响应略慢但稳定性突出，可作为展示性能上限。"],
            ["速度/稳定平衡组合", "寻找综合性能最优点。", "3 个", "tc ≤1.2 s；tb ≤0.8 s；8000 cycles ≥90%。", "优先作为后续论文或报告主路线。"],
        ],
        [1700, 2200, 1100, 2600, 1760],
        font_size=8.7,
        center_cols=[2],
    )
    add_table(
        doc,
        ["cycle", "必测内容", "备注/分析"],
        [
            ["1", "初始 CV、ΔT、tc/tb、λmax", "建立真实初始状态。"],
            ["50", "ΔT、残余峰", "观察早期快速衰减。"],
            ["100", "ΔT、残余峰", "若保持率 <70%，优先淘汰条件。"],
            ["300", "ΔT", "补充衰减曲线。"],
            ["500", "ΔT、CV", "判断是否进入稳定平台。"],
            ["1000", "ΔT、tc/tb", "观察响应是否随循环变化。"],
            ["3000", "ΔT", "中期稳定性。"],
            ["5000", "ΔT、CV", "检查残余峰和可逆性。"],
            ["8000", "ΔT、tc/tb", "若保持率 ≥85%，继续 13000 cycles。"],
            ["13000", "ΔT、CV、tc/tb", "长循环验证终点。"],
        ],
        [1000, 3600, 4760],
        font_size=8.7,
        center_cols=[0],
    )

    add_heading(doc, "8. 封装与测试统一要求", 1)
    add_table(
        doc,
        ["类别", "统一要求", "目的", "记录项"],
        [
            ["UV 固化", "黑色胶带遮挡变色层；UV 胶只在边缘或远离电解质处少量点写。", "避免 UV 改变变色层或污染电解质。", "遮挡方式、曝光时间、胶位置。"],
            ["凝胶接触", "记录热压温度、压力、时间，确认凝胶完全覆盖活性区。", "减少界面接触差异造成的响应/循环波动。", "温度、压力、时间、气泡/干区。"],
            ["边缘密封", "液态对照重点验证沙林膜-玻璃边缘全封闭；凝胶器件记录边缘泄漏。", "把封装失效与材料失效区分开。", "封装前后和循环后照片。"],
            ["响应时间", "统一采用达到 90% ΔT 的时间，光谱采集间隔 0.1-0.2 s。", "保证不同器件数据可比。", "λmax、ΔT、tc/tb、采样间隔。"],
            ["光谱测试", "重新放置器件后必须重新扣背景，并记录夹具位置和样品角度。", "避免光路变化导致 ΔT 异常。", "背景、角度、夹具位置。"],
            ["循环前后", "CA 循环前后均做 CV。", "判断残余峰与可逆性。", "CV 峰位、峰面积、残余吸收。"],
        ],
        [1300, 3300, 2500, 2260],
        font_size=8.5,
    )

    add_heading(doc, "9. 一周执行安排", 1)
    add_table(
        doc,
        ["时间", "工作内容", "输出", "决策点"],
        [
            ["Day 1", "同批 TiO2 基底；制备 D-100、D-5x、D-10x、D-5x-rinse。", "染浴体积、液面覆盖照片、DCM 冲洗记录。", "确认染浴覆盖是否可重复。"],
            ["Day 2", "制备 ATO-1、ATO-2、ATO-3；组装关键组合。", "D-100/ATO-2、D-5x/ATO-2、D-5x-rinse/ATO-2、D-5x/ATO-1、D-5x/ATO-3。", "确认器件外观、气泡和重叠面积。"],
            ["Day 3", "初始 CV/CA/光谱；比较 -2.0 V/0 V 与 -2.0 V/+0.3 V。", "ΔT、tc/tb、残余峰。", "判断是否需要正压褪色。"],
            ["Day 4-5", "500 cycles 快速筛选。", "1、50、100、300、500 cycles ΔT。", "前 100 cycles 保持率 <70% 的条件优先淘汰。"],
            ["Day 6-7", "最优 2 条件重复制备或继续 1000-3000 cycles。", "重复性与中期稳定性数据。", "决定下一周进入 TiO2、ATO 或 PBI 微调。"],
        ],
        [1100, 3300, 3000, 1960],
        font_size=8.3,
    )

    add_heading(doc, "10. 决策树", 1)
    add_table(
        doc,
        ["观察结果", "判断", "下一步动作"],
        [
            ["D-5x/D-10x 明显优于 D-100，残余峰不严重。", "PBI 有效负载是主要变量，且过量吸附风险可控。", "进入 TiO2 厚度优化。"],
            ["D-5x/D-10x 对比度高但残余峰严重。", "负载过高或聚集/氧化不完全。", "优先优化 DCM 冲洗、PBI 浓度/时间和褪色正电压。"],
            ["ATO-3 明显改善褪色和循环保持率。", "对电极容量或面积匹配限制明显。", "进入 ATO 层数/退火或面积匹配优化。"],
            ["ATO 层数影响小，但 +0.3 V 褪色明显改善。", "电压窗口是主控因素之一。", "确定温和正压褪色窗口，避免 +0.5 V 过度氧化。"],
            ["所有条件前 500 cycles 都快速衰减。", "封装、凝胶接触或 UV 暴露可能主导失效。", "优先检查封装和界面接触，不继续增加 PBI 负载。"],
        ],
        [3300, 3000, 3060],
        font_size=8.5,
    )

    add_heading(doc, "11. 数据记录模板", 1)
    add_heading(doc, "11.1 电极制备记录", 2)
    add_table(
        doc,
        ["样品", "TiO2层数", "TiO2退火", "PBI浓度(mM)", "染浴体积", "染浴时间", "DCM冲洗", "后处理", "电极颜色/吸收"],
        [["", "", "", "", "", "", "", "", ""] for _ in range(4)],
        [850, 900, 950, 1100, 1000, 1000, 1000, 900, 1660],
        font_size=8.2,
        center_cols=[0, 1, 2, 3, 4, 5, 6],
    )
    add_heading(doc, "11.2 器件性能记录", 2)
    add_table(
        doc,
        ["样品", "ATO层数", "ATO退火", "面积比", "电压窗口", "ΔT(λmax)", "λmax", "tc", "tb", "500cy保持率", "残余峰"],
        [["", "", "", "", "", "", "", "", "", "", ""] for _ in range(4)],
        [750, 800, 850, 750, 1100, 1000, 700, 650, 650, 1100, 1010],
        font_size=8.0,
        center_cols=list(range(0, 11)),
    )
    add_heading(doc, "11.3 电荷匹配记录", 2)
    add_table(
        doc,
        ["样品", "Q_PBI (mC/cm²)", "Q_ATO (mC/cm²)", "Q_ATO/Q_PBI", "褪色是否完全", "判断"],
        [["", "", "", "", "", ""] for _ in range(4)],
        [1000, 1800, 1800, 1600, 1700, 1460],
        font_size=8.3,
        center_cols=[0, 1, 2, 3, 4],
    )
    add_heading(doc, "11.4 循环稳定性记录", 2)
    add_table(
        doc,
        ["样品", "初始ΔT", "50cy", "100cy", "300cy", "500cy", "1000cy", "3000cy", "5000cy", "8000cy", "13000cy", "失效现象"],
        [["", "", "", "", "", "", "", "", "", "", "", ""] for _ in range(4)],
        [700, 800, 650, 700, 700, 700, 800, 800, 800, 800, 850, 1060],
        font_size=7.8,
        center_cols=list(range(0, 11)),
    )

    add_heading(doc, "12. 参考文献", 1)
    refs = [
        "H. J. Kim, J. K. Seo, Y.-J. Kim et al. Formation of ultrafast-switching viologen-anchored TiO2 electrochromic device. Solar Energy Materials and Solar Cells, 2009, 93, 1982-1987. DOI: 10.1016/j.solmat.2009.05.007.",
        "D. Nunes, T. L. Freire, A. Barranger et al. TiO2 Nanostructured Films for Electrochromic Paper Based-Devices. Applied Sciences, 2020, 10, 1200. DOI: 10.3390/app10041200.",
        "R. Lundy, E. R. Draper, J. J. Walsh et al. Amino acid appended perylene bisimides: self-assembly, immobilization on nanocrystalline TiO2, and electrochromic properties. New Journal of Chemistry, 2018, 42, 19020-19025. DOI: 10.1039/C8NJ04214D.",
        "M. Gratzel. Dye-sensitized solar cells. Journal of Photochemistry and Photobiology C, 2003, 4, 145-153. DOI: 10.1016/S1389-5567(03)00026-1.",
        "F. Sauvage, J.-D. Decoppet, M. Zhang, S. M. Zakeeruddin, P. Comte, M. K. Nazeeruddin, P. Wang, M. Gratzel. Effect of Sensitizer Adsorption Temperature on the Performance of Dye-Sensitized Solar Cells. Journal of the American Chemical Society, 2011, 133, 9304-9310. DOI: 10.1021/ja110541t.",
        "D. Choi, M. Lee, H. Kim et al. Investigation of dry-deposited ion storage layers using various oxide particles to enhance electrochemical performance. Solar Energy Materials and Solar Cells, 2017, 179, 422-428. DOI: 10.1016/j.solmat.2017.10.001.",
        "S. Lee et al. Polymer-Assisted Generation of Antimony-Doped SnO2 Nanoparticles with High Conductivity. Small, 2008, 4, 1906-1912.",
        "Y. Alesanco, A. Vinuales, J. Rodriguez et al. All-in-One Gel-Based Electrochromic Devices: Strengths and Recent Developments. Materials, 2018, 11, 414. DOI: 10.3390/ma11030414.",
        "P.-W. Chen, C.-T. Chang, T.-F. Ko et al. Fast response of complementary electrochromic device based on WO3/NiO electrodes. Scientific Reports, 2020, 10, 8431. DOI: 10.1038/s41598-020-65191-x.",
        "R. Li, X. Ma, J. Li et al. Flexible and high-performance electrochromic devices enabled by self-assembled 2D TiO2/MXene heterostructures. Nature Communications, 2021, 12, 1587. DOI: 10.1038/s41467-021-21852-7.",
        "V. K. Thakur, G. Ding, J. Ma, P. S. Lee, X. Lu. Hybrid Materials and Polymer Electrolytes for Electrochromic Device Applications. Advanced Materials, 2012, 24, 4071-4096. DOI: 10.1002/adma.201200213.",
        "P. M. S. Monk, R. J. Mortimer, D. R. Rosseinsky. Electrochromism and Electrochromic Devices. Cambridge University Press, 2007.",
    ]
    for i, ref in enumerate(refs, 1):
        add_paragraph(doc, f"{i}. {ref}", size=8.7, after=3, line=1.15)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    path = build_doc()
    print(path)
